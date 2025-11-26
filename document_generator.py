"""
Генератор DOCX документов в корпоративном стиле ForteBank
"""
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class CorporateDocxGenerator:
    """Генератор DOCX документов в корпоративном стиле"""
    
    # ForteBank Corporate Colors
    FORTE_BLUE = RGBColor(0, 82, 155)      # #00529B
    FORTE_DARK = RGBColor(34, 34, 34)      # #222222
    FORTE_GRAY = RGBColor(102, 102, 102)   # #666666
    FORTE_LIGHT_GRAY = RGBColor(242, 242, 242)  # #F2F2F2
    
    def __init__(self, output_dir: str = "docs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_docx(
        self,
        markdown_content: str,
        doc_type: str,
        session_id: str,
        user_title: Optional[str] = None
    ) -> str:
        """
        Генерирует DOCX документ из Markdown контента.
        
        Args:
            markdown_content: Контент в Markdown формате
            doc_type: Тип документа
            session_id: ID сессии
            user_title: Название документа
        
        Returns:
            Путь к созданному файлу
        """
        doc = Document()
        
        # Применяем корпоративные стили
        self._apply_corporate_styles(doc)
        
        # Добавляем корпоративную шапку
        self._add_corporate_header(doc, doc_type)
        
        # Парсим и добавляем контент
        self._parse_and_add_content(doc, markdown_content)
        
        # Добавляем корпоративный футер
        self._add_corporate_footer(doc)
        
        # Генерируем имя файла
        filepath = self._generate_filepath(doc_type, session_id, user_title)
        
        # Сохраняем
        doc.save(filepath)
        logger.info(f"DOCX saved: {filepath}")
        
        return str(filepath)
    
    def _apply_corporate_styles(self, doc: Document):
        """Применяет корпоративные стили к документу"""
        
        # Настройка стилей
        styles = doc.styles
        
        # Стиль для заголовка H1
        if 'Corporate Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Corporate Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Arial'
            h1_font.size = Pt(24)
            h1_font.bold = True
            h1_font.color.rgb = self.FORTE_BLUE
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        
        # Стиль для заголовка H2
        if 'Corporate Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Corporate Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Arial'
            h2_font.size = Pt(18)
            h2_font.bold = True
            h2_font.color.rgb = self.FORTE_DARK
            h2_style.paragraph_format.space_before = Pt(10)
            h2_style.paragraph_format.space_after = Pt(4)
        
        # Стиль для заголовка H3
        if 'Corporate Heading 3' not in [s.name for s in styles]:
            h3_style = styles.add_style('Corporate Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            h3_font = h3_style.font
            h3_font.name = 'Arial'
            h3_font.size = Pt(14)
            h3_font.bold = True
            h3_font.color.rgb = self.FORTE_GRAY
            h3_style.paragraph_format.space_before = Pt(8)
            h3_style.paragraph_format.space_after = Pt(3)
        
        # Стиль для обычного текста
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)
        normal.font.color.rgb = self.FORTE_DARK
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(6)
    
    def _add_corporate_header(self, doc: Document, doc_type: str):
        """Добавляет корпоративную шапку документа"""
        
        # Логотип (текстовый, так как у нас нет файла)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('ForteBank')
        run.font.name = 'Arial'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.FORTE_BLUE
        
        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Business Analysis Department')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = self.FORTE_GRAY
        
        # Разделитель
        doc.add_paragraph('_' * 80)
        
        # Метаинформация
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        
        # Настройка ширины колонок
        meta_table.columns[0].width = Inches(2.0)
        meta_table.columns[1].width = Inches(4.5)
        
        # Заполнение метаданных
        doc_type_names = {
            "new_feature": "Business Requirements Document",
            "bug_fix": "Bug Fix Requirements",
            "process_change": "Process Change Request",
            "integration": "Integration Requirements",
            "data_request": "Data Request Specification"
        }
        
        meta_data = [
            ("Тип документа:", doc_type_names.get(doc_type, "Document")),
            ("Дата создания:", datetime.now().strftime("%d.%m.%Y %H:%M")),
            ("Автор:", "AI Business Analyst"),
            ("Статус:", "Draft - На согласовании")
        ]
        
        for i, (key, value) in enumerate(meta_data):
            row = meta_table.rows[i]
            
            # Ключ
            key_cell = row.cells[0]
            key_run = key_cell.paragraphs[0].add_run(key)
            key_run.font.bold = True
            key_run.font.size = Pt(10)
            key_run.font.color.rgb = self.FORTE_GRAY
            
            # Значение
            val_cell = row.cells[1]
            val_run = val_cell.paragraphs[0].add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = self.FORTE_DARK
        
        # Пустая строка
        doc.add_paragraph()
    
    def _parse_and_add_content(self, doc: Document, markdown_content: str):
        """Парсит Markdown и добавляет контент в документ"""
        
        lines = markdown_content.split('\n')
        in_code_block = False
        in_table = False
        table_data = []
        code_lines = []
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Code block
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    in_code_block = False
                    if code_lines:
                        self._add_code_block(doc, '\n'.join(code_lines))
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # Table
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # Skip separator row
                if all(c.strip().replace('-', '').replace(':', '') == '' for c in cells):
                    i += 1
                    continue
                
                table_data.append(cells)
                i += 1
                
                # Check if next line is still table
                if i < len(lines) and '|' not in lines[i]:
                    in_table = False
                    self._add_table(doc, table_data)
                    table_data = []
                continue
            
            # Heading 1
            if line.startswith('# ') and not line.startswith('## '):
                heading = doc.add_paragraph(line[2:], style='Corporate Heading 1')
                i += 1
                continue
            
            # Heading 2
            if line.startswith('## ') and not line.startswith('### '):
                heading = doc.add_paragraph(line[3:], style='Corporate Heading 2')
                i += 1
                continue
            
            # Heading 3
            if line.startswith('### '):
                heading = doc.add_paragraph(line[4:], style='Corporate Heading 3')
                i += 1
                continue
            
            # Horizontal rule
            if line.strip() == '---' or line.strip().startswith('---'):
                doc.add_paragraph('_' * 80)
                i += 1
                continue
            
            # Bullet list
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(text, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                i += 1
                continue
            
            # Numbered list
            if re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(text, style='List Number')
                p.paragraph_format.left_indent = Inches(0.25)
                i += 1
                continue
            
            # Bold text inline
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    if j % 2 == 1:  # Odd indices are bold
                        run.bold = True
                i += 1
                continue
            
            # Empty line
            if not line.strip():
                i += 1
                continue
            
            # Normal paragraph
            if line.strip():
                # Удаляем markdown форматирование
                clean_text = line.strip()
                clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)  # Bold
                clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)      # Italic
                clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)        # Code
                
                doc.add_paragraph(clean_text)
            
            i += 1
    
    def _add_code_block(self, doc: Document, code: str):
        """Добавляет блок кода"""
        p = doc.add_paragraph()
        p.style = 'Normal'
        
        # Фон для кода
        shading_elm = p._element.get_or_add_pPr()
        
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_table(self, doc: Document, table_data: list):
        """Добавляет таблицу в документ"""
        if not table_data:
            return
        
        # Создаем таблицу
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Заполняем данные
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                
                # Первая строка - заголовок
                if i == 0:
                    run = cell.paragraphs[0].add_run(cell_data)
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.FORTE_BLUE
                else:
                    run = cell.paragraphs[0].add_run(cell_data)
                    run.font.size = Pt(10)
        
        # Добавляем отступ после таблицы
        doc.add_paragraph()
    
    def _add_corporate_footer(self, doc: Document):
        """Добавляет корпоративный футер"""
        
        # Разделитель
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        
        # Футер
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = (
            "ForteBank JSC | Almaty, Kazakhstan\n"
            "Business Analysis Department\n"
            f"Document generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        )
        
        run = footer.add_run(footer_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = self.FORTE_GRAY
        run.font.italic = True
    
    def _generate_filepath(
        self,
        doc_type: str,
        session_id: str,
        user_title: Optional[str]
    ) -> Path:
        """Генерирует путь к файлу"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if user_title:
            safe_title = self._sanitize_filename(user_title)
            filename = f"{timestamp}_{doc_type}_{safe_title}.docx"
        else:
            filename = f"{timestamp}_{doc_type}_{session_id}.docx"
        
        return self.output_dir / filename
    
    def _sanitize_filename(self, title: str) -> str:
        """Очищает строку для имени файла"""
        safe = re.sub(r'[<>:"/\\|?*]', '', title)
        safe = safe.replace(' ', '_')
        safe = safe[:50]
        return safe.lower()
    
    def list_documents(self) -> list:
        """Список всех созданных документов"""
        docs = []
        for file in self.output_dir.glob("*.docx"):
            # Пропускаем временные файлы Word
            if file.name.startswith('~$'):
                continue
            
            docs.append({
                "filename": file.name,
                "path": str(file),
                "created": datetime.fromtimestamp(file.stat().st_mtime),
                "size": file.stat().st_size
            })
        
        docs.sort(key=lambda x: x["created"], reverse=True)
        return docs