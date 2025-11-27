"""
Mermaid → PNG Converter
Конвертация Mermaid диаграмм в PNG для вставки в DOCX
"""
import logging
import base64
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class MermaidToPNGConverter:
    """Конвертер Mermaid диаграмм в PNG"""

    def __init__(self, use_mermaid_cli: bool = True):
        """
        Args:
            use_mermaid_cli: Использовать mermaid-cli (требует npm install -g @mermaid-js/mermaid-cli)
        """
        self.use_cli = use_mermaid_cli

        if use_mermaid_cli:
            self._check_mermaid_cli()

    def _check_mermaid_cli(self):
        """Проверить наличие mermaid-cli"""
        try:
            result = subprocess.run(
                ["mmdc", "--version"],
                capture_output=True,
                text=True,
                timeout=5
            )
            logger.info(f"✅ mermaid-cli found: {result.stdout.strip()}")
        except FileNotFoundError:
            logger.warning("⚠️ mermaid-cli not found. Install: npm install -g @mermaid-js/mermaid-cli")
        except Exception as e:
            logger.warning(f"⚠️ mermaid-cli check failed: {e}")

    def convert_to_png(
            self,
            mermaid_code: str,
            output_path: Optional[str] = None,
            width: int = 800,
            background: str = "white"
    ) -> str:
        """
        Конвертировать Mermaid в PNG.

        Args:
            mermaid_code: Mermaid код
            output_path: Путь для сохранения PNG (если None - создается временный)
            width: Ширина изображения
            background: Цвет фона

        Returns:
            Путь к созданному PNG файлу
        """
        if self.use_cli:
            return self._convert_with_cli(mermaid_code, output_path, width, background)
        else:
            return self._convert_with_api(mermaid_code, output_path)

    def _convert_with_cli(
            self,
            mermaid_code: str,
            output_path: Optional[str],
            width: int,
            background: str
    ) -> str:
        """Конвертация через mermaid-cli (mmdc)"""
        # Создаем временный файл с Mermaid кодом
        with tempfile.NamedTemporaryFile(
                mode='w',
                suffix='.mmd',
                delete=False,
                encoding='utf-8'
        ) as tmp_input:
            tmp_input.write(mermaid_code)
            input_path = tmp_input.name

        # Определяем output path
        if output_path is None:
            output_path = tempfile.mktemp(suffix='.png')

        try:
            # Запускаем mmdc
            cmd = [
                "mmdc",
                "-i", input_path,
                "-o", output_path,
                "-w", str(width),
                "-b", background,
                "-t", "default"
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode == 0:
                logger.info(f"✅ Mermaid → PNG: {output_path}")
                return output_path
            else:
                logger.error(f"❌ mmdc failed: {result.stderr}")
                raise Exception(f"mmdc conversion failed: {result.stderr}")

        finally:
            # Удаляем временный input файл
            Path(input_path).unlink(missing_ok=True)

    def _convert_with_api(
            self,
            mermaid_code: str,
            output_path: Optional[str]
    ) -> str:
        """
        Конвертация через онлайн API (fallback).
        Использует https://mermaid.ink/
        """
        import requests

        # Кодируем Mermaid в base64
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')

        # URL для mermaid.ink API
        url = f"https://mermaid.ink/img/{encoded}"

        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Определяем output path
            if output_path is None:
                output_path = tempfile.mktemp(suffix='.png')

            # Сохраняем PNG
            with open(output_path, 'wb') as f:
                f.write(response.content)

            logger.info(f"✅ Mermaid → PNG (API): {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"❌ API conversion failed: {e}")
            raise

    def convert_and_embed_base64(
            self,
            mermaid_code: str,
            width: int = 800
    ) -> str:
        """
        Конвертировать и вернуть как base64 строку.
        Удобно для вставки в HTML/JSON.

        Returns:
            Base64 строка: "data:image/png;base64,iVBORw0KGgo..."
        """
        png_path = self.convert_to_png(mermaid_code, width=width)

        try:
            with open(png_path, 'rb') as f:
                png_data = f.read()

            base64_str = base64.b64encode(png_data).decode('utf-8')
            return f"data:image/png;base64,{base64_str}"

        finally:
            # Удаляем временный файл
            Path(png_path).unlink(missing_ok=True)


# ============================================================================
# DOCX IMAGE INSERTER
# ============================================================================

class DOCXImageInserter:
    """Вставка PNG изображений в DOCX документы"""

    @staticmethod
    def insert_image_after_paragraph(
            doc,
            paragraph_text: str,
            image_path: str,
            width_inches: float = 6.0
    ):
        """
        Вставить изображение после параграфа с определенным текстом.

        Args:
            doc: python-docx Document объект
            paragraph_text: Текст параграфа после которого вставить
            image_path: Путь к PNG
            width_inches: Ширина в дюймах
        """
        from docx.shared import Inches

        for i, para in enumerate(doc.paragraphs):
            if paragraph_text in para.text:
                # Вставляем новый параграф после текущего
                new_para = doc.paragraphs[i]._element
                parent = new_para.getparent()

                # Создаем новый параграф
                from docx.oxml import OxmlElement
                new_p = OxmlElement('w:p')
                parent.insert(parent.index(new_para) + 1, new_p)

                # Добавляем изображение
                from docx.text.paragraph import Paragraph
                p = Paragraph(new_p, doc)
                run = p.add_run()
                run.add_picture(image_path, width=Inches(width_inches))

                logger.info(f"✅ Image inserted after: {paragraph_text[:50]}...")
                return True

        logger.warning(f"⚠️ Paragraph not found: {paragraph_text[:50]}...")
        return False

    @staticmethod
    def append_image(
            doc,
            image_path: str,
            width_inches: float = 6.0,
            caption: Optional[str] = None
    ):
        """
        Добавить изображение в конец документа.

        Args:
            doc: python-docx Document
            image_path: Путь к изображению
            width_inches: Ширина
            caption: Подпись под изображением
        """
        from docx.shared import Inches

        # Добавляем параграф с изображением
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        # Добавляем подпись если есть
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.add_run(caption).italic = True
            caption_para.alignment = 1  # Center

        logger.info(f"✅ Image appended: {image_path}")


# ============================================================================
# COMBINED WORKFLOW
# ============================================================================

def mermaid_to_docx_workflow(
        mermaid_code: str,
        doc,
        insert_after: Optional[str] = None,
        width_inches: float = 6.0,
        caption: Optional[str] = None
):
    """
    Полный workflow: Mermaid → PNG → вставка в DOCX.

    Args:
        mermaid_code: Mermaid диаграмма
        doc: python-docx Document
        insert_after: Текст параграфа после которого вставить (если None - в конец)
        width_inches: Ширина изображения
        caption: Подпись

    Returns:
        True если успешно
    """
    try:
        # Конвертируем Mermaid → PNG
        converter = MermaidToPNGConverter(use_mermaid_cli=True)
        png_path = converter.convert_to_png(mermaid_code)

        # Вставляем в DOCX
        inserter = DOCXImageInserter()

        if insert_after:
            success = inserter.insert_image_after_paragraph(
                doc, insert_after, png_path, width_inches
            )
        else:
            inserter.append_image(doc, png_path, width_inches, caption)
            success = True

        # Удаляем временный PNG
        Path(png_path).unlink(missing_ok=True)

        return success

    except Exception as e:
        logger.error(f"❌ Mermaid → DOCX workflow failed: {e}")
        return False


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    # Тест 1: Mermaid → PNG через CLI
    converter = MermaidToPNGConverter(use_mermaid_cli=True)

    mermaid_code = """
graph TD
    A[Start] --> B[Process]
    B --> C{Decision?}
    C -->|Yes| D[End]
    C -->|No| B
"""

    png_path = converter.convert_to_png(mermaid_code, "test_diagram.png")
    print(f"Created: {png_path}")

    # Тест 2: Base64
    base64_img = converter.convert_and_embed_base64(mermaid_code)
    print(f"Base64 length: {len(base64_img)}")